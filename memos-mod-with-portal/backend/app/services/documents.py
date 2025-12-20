from __future__ import annotations
from typing import Optional, List, Dict
from pathlib import Path
from .files import ensure_dir
from ..core.config import TEMPLATE_PATH
import os
import shutil
import subprocess
import tempfile
from docx.shared import Pt  # type: ignore
from docx.oxml.ns import qn  # type: ignore

# Attempt to import python-docx optionally
try:
    from docx import Document  # type: ignore
    from docx.shared import Inches
except Exception:
    Document = None  # type: ignore
    Inches = None  # type: ignore

def _replace_in_paragraph(par, mapping: Dict[str, str]):
    for key, val in mapping.items():
        if key in par.text:
            par.text = par.text.replace(key, val)
    
    
            # Aplicar Arial 10 a todo el párrafo
            for run in par.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)

                # Para que Word respete la fuente
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:ascii'), 'Arial')
                rFonts.set(qn('w:hAnsi'), 'Arial')
                rFonts.set(qn('w:eastAsia'), 'Arial')
                rFonts.set(qn('w:cs'), 'Arial')



def _replace_in_table(table, mapping: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, mapping)

def generate_doc_from_template(ctx: Dict[str,str], evid_paths: List[Path], out_docx: Path):
    ensure_dir(out_docx)
    if Document is None or not TEMPLATE_PATH.exists():
        # Fallback: simple text placeholder if docx not possible
        out_txt = out_docx.with_suffix(".txt")
        out_txt.write_text(f"[PLACEHOLDER] Falta plantilla o python-docx.\n{ctx}", encoding="utf-8")
        return out_txt  # may not be .docx
    # Build from real .docx template
    doc = Document(str(TEMPLATE_PATH))
    mapping = {f"{{{{{k}}}}}": str(v) for k, v in ctx.items()}
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in doc.tables:
        _replace_in_table(t, mapping)
    # Evidences
    if evid_paths and Inches is not None:
        doc.add_page_break()
        title = doc.add_paragraph()
        title.alignment = 1  # 0 left, 1 center, 2 right, 3 justify
        run = title.add_run("ANEXOS")
        run.bold = True

        doc.add_paragraph("")  # espacio
        #doc.add_paragraph("Evidencias:")
        #doc.add_paragraph("")  # espacio

        for img in evid_paths:
            try:
                doc.add_picture(str(img), width=Inches(4.0))  # type: ignore
            except Exception:
                pass
    doc.save(out_docx)
    return out_docx

class DocxToPdfError(RuntimeError):
    pass

def docx_to_pdf(
    input_path: str | os.PathLike,
    output_dir: str | os.PathLike | None = None,
    timeout: int = 120,
) -> Path:
    """
    Convierte un archivo .docx a .pdf usando LibreOffice headless.
    Devuelve la ruta absoluta del PDF generado.
    Lanza DocxToPdfError/ValueError/FileNotFoundError en caso de error.
    """

    print("[docx_to_pdf] Convirtiendo a PDF:", input_path)

    in_path = Path(input_path).expanduser().resolve()
    if not in_path.exists():
        raise FileNotFoundError(f"No existe: {in_path}")
    if in_path.suffix.lower() != ".docx":
        raise ValueError("El archivo de entrada debe ser .docx")

    # Buscar LibreOffice
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise DocxToPdfError(
            "LibreOffice no está instalado o no está en PATH. "
            "Instálalo con: sudo apt-get update && sudo apt-get install -y libreoffice"
        )

    out_dir = Path(output_dir).expanduser().resolve() if output_dir else in_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    out_pdf = out_dir / (in_path.stem + ".pdf")

    # HOME temporal para evitar problemas de primera ejecución
    with tempfile.TemporaryDirectory(prefix="lo-home-") as lo_home:
        env = os.environ.copy()
        env["HOME"] = lo_home
        env.setdefault("LANG", "en_US.UTF-8")

        cmd = [
            soffice,
            "--headless",
            "--nologo",
            "--nolockcheck",
            "--nodefault",
            "--norestore",
            "--invisible",
            "--convert-to", "pdf:writer_pdf_Export",
            "--outdir", str(out_dir),
            str(in_path),
        ]

        try:
            proc = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=env,
                timeout=timeout,
                check=False,
                text=True,
            )
        except subprocess.TimeoutExpired:
            raise DocxToPdfError(f"Conversión tardó más de {timeout}s y fue cancelada.")

        if proc.returncode != 0 or not out_pdf.exists():
            msg = (proc.stderr or "") + "\n" + (proc.stdout or "")
            raise DocxToPdfError(f"LibreOffice no pudo convertir el archivo.\n{msg}")

    return out_pdf.resolve()


def try_export_pdf(path_docx: Path) -> Optional[Path]:
    """
    Wrapper compatible con la versión antigua:
    - Recibe Path al .docx
    - Intenta generar PDF en mismo directorio
    - Devuelve Path al PDF o None si falla (sin lanzar excepción)
    """
    print("[try_export_pdf] Intentando exportar PDF de:", path_docx)
    try:
        if path_docx.suffix.lower() != ".docx":
            print("[try_export_pdf] No es un .docx válido.")
            return None
        pdf_path = docx_to_pdf(path_docx, output_dir=path_docx.parent)
        return pdf_path
    except Exception:
        print("[try_export_pdf] Falló la conversión a PDF.")
        return None


#def try_export_pdf(path_docx: Path) -> Optional[Path]:
#    if docx2pdf_convert is None:
#        return None
#    out_pdf = path_docx.with_suffix(".pdf")
#    if pythoncom:
#        pythoncom.CoInitialize()
#    try:
#       docx2pdf_convert(str(path_docx), str(out_pdf))
#       return out_pdf
#    except Exception:
#        return None
#    finally:
#        if pythoncom:
#            pythoncom.CoUninitialize()