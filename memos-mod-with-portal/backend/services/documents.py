from typing import Optional, List, Dict
from pathlib import Path
from .files import ensure_dir
from ..core.config import TEMPLATE_PATH
from ..core.compat import docx2pdf_convert, pythoncom

# Attempt to import python-docx optionally
try:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore
except Exception:
    Document = None  # type: ignore
    Inches = None  # type: ignore


import subprocess
from pathlib import Path

def try_export_pdf_linux(path_docx: Path) -> Path | None:
    out_pdf = path_docx.with_suffix(".pdf")
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(path_docx.parent), str(path_docx)],
            check=True
        )
        return out_pdf
    except Exception as e:
        print("Error convirtiendo:", e)
        return None


def _replace_in_paragraph(par, mapping: Dict[str, str]):
    for key, val in mapping.items():
        if key in par.text:
            par.text = par.text.replace(key, val)

def _replace_in_table(table, mapping: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, mapping)

def generate_doc_from_template(ctx: Dict[str,str], evid_paths: List[Path], out_docx: Path):
    ensure_dir(out_docx)
    if Document is None or not TEMPLATE_PATH.exists():
        # Fallback: simple text placeholder if docx not possible
        out_txt = out_docx.with_suffix(".txt")
        out_txt.write_text(f"[PLACEHOLDER] Falta plantilla o python-docx.\n{ctx}", encoding="utf-8")
        return out_txt  # may not be .docx
    # Build from real .docx template
    doc = Document(str(TEMPLATE_PATH))
    mapping = {f"{{{{{k}}}}}": str(v) for k, v in ctx.items()}
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in doc.tables:
        _replace_in_table(t, mapping)
    # Evidences
    if evid_paths and Inches is not None:
        doc.add_page_break()
        doc.add_paragraph("Evidencias:")
        for img in evid_paths:
            try:
                doc.add_picture(str(img), width=Inches(6.0))  # type: ignore
            except Exception:
                pass
    doc.save(out_docx)
    return out_docx